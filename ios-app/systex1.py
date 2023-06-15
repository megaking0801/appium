from appium.webdriver.common.appiumby import AppiumBy
from compare import CompareName , OddLot ,Leave ,KLine , CompareNumber, ClickBuyAndSell
from excelTest import PrintExcel
from word import CompanyName
from time import sleep
from selenium.webdriver.common.by import By
from appium.webdriver.common.touch_action import TouchAction
from appium import webdriver
from appium.webdriver.common.mobileby import MobileBy
import re
from selenium.common.exceptions import NoSuchElementException
from docx import Document
from docx.shared import Cm,Pt
import datetime

doc = Document('test.docx')

class SearchSystex1:
    def __init__(driver):
        try:
            systex = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠 6214")

            systex.click()

            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            el1.click()
            sleep(3)
            el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            el2.click()
            sleep(3)
            el3 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el3.click()
            sleep(3)
        

            
            # 找到所有的 XCUIElementTypeStaticText 元素
            # 找到所有 type 為 XCUIElementTypeStaticText 的元素
            elements = driver.find_elements(By.XPATH, "//XCUIElementTypeStaticText")

            for element in elements:
    # 取得元件的名稱和 value 屬性
                element_name = element.get_attribute("name")
                element_value = element.get_attribute("value")
    
    # 判斷 value 屬性是否包含 "-"
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

            a = table.cell(0,0).paragraphs[0].add_run("測試內容：判斷 value 屬性是否包含-" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            
            if "-" in element_value:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + f"{element_name}" + "數值錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print(f"{element_name}: false")
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "數值無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("沒有元件value有-")

            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()

            doc.save('test.docx')
        except Exception:
            print("exception")
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：判斷 value 屬性是否包含-" )
            a.font.name='黑體'
            a.font.size=Pt(16)

            b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            

            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 
            
            doc.add_page_break()

            doc.save('test.docx')
            
            
        
        
        
        
        
        #進入零股畫面，確認零股名稱正確
        try:
            sleep(3)
            enterOddLot = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeStaticText[@name=\"零股\"]") #零股
            sleep(3)
            enterOddLot.click()

            oddLot = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠 6214")
            oddLot.click
            oddLotName = oddLot.text
            print(oddLotName)

            
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

            a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認零股名稱是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            
            if(oddLotName!="精誠 6214"):
                b = table.cell(0,1).paragraphs[0].add_run("零股名稱錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("零股名稱錯誤")
            else:
                b = table.cell(0,1).paragraphs[0].add_run("零股名稱無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("零股名稱無誤")

            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()

            doc.save('test.docx')
        except Exception:
            print("exception")
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認零股名稱是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)

            b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            

            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 
            
            doc.add_page_break()

            doc.save('test.docx')
        #離開
        el9 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"零股\"]")
        sleep(3)
        el9.click()










#檢查K線名稱
        try:
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認K線名稱是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)

            try:
                el2 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeStaticText[@name=\"K線\"]")
                el2.click()

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)

                el7 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
                sleep(3)
                el7.click()
                sleep(3)
                el8 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
                sleep(3)
                el8.click()
                sleep(3)
                print("有說明書")
                b = table.cell(0,1).paragraphs[0].add_run("有說明書 ")
                b.font.name='黑體'
                b.font.size=Pt(16)
                doc.add_page_break()
                doc.save('test.docx')
            except Exception:
                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                print("沒有說明書")


            getName = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeApplication[@name=\"樹精靈(測)\"]/XCUIElementTypeWindow[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[1]/XCUIElementTypeOther/XCUIElementTypeOther[1]")
            kLineName = (getName.text +" ").strip()
            if(kLineName != "6214精誠"):
                print("K線不相同")
                b = table.cell(0,1).paragraphs[0].add_run("K線不相同" )
                b.font.name='黑體'
                b.font.size=Pt(16)            
            else:
                b = table.cell(0,1).paragraphs[0].add_run("K線相同" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("K線相同")

            doc.add_page_break()
            doc.save('test.docx')

        except Exception:
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認K線名稱是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9)
            print("error")
            doc.add_page_break()
            doc.save('test.docx')