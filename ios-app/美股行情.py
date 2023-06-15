from appium.webdriver.common.appiumby import AppiumBy
from time import sleep
from docx import Document
from appium.webdriver.common.mobileby import MobileBy
from appium.webdriver.common.touch_action import TouchAction
from selenium.webdriver.common.action_chains import ActionChains,ActionBuilder
from selenium.webdriver.common.actions.pointer_input import PointerInput
from selenium.webdriver.common.actions import interaction
from docx import Document
from openpyxl import load_workbook
from docx.shared import Cm,Pt
from datetime import datetime
from openpyxl import Workbook

class a_stock:
    def __init__(driver):

        doc = Document('test.docx')
        wb = load_workbook('test.xlsx')
        ws = wb.active
        el1 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"自選庫存\"]")
        sleep(3)
        el1.click()

        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(237, 639)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.pause(0.1)
        actions.w3c_actions.pointer_action.release()
        actions.perform()

        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(24, 435)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.move_to_location(24, 543)
        actions.w3c_actions.pointer_action.release()
        actions.perform()
    

        el5 = driver.find_element(by=AppiumBy.XPATH, value="(//XCUIElementTypeStaticText[@name=\"美股專區\"])[2]")
        el5.click()
        sleep(10)
        el6 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="科技龍頭")
        b1 = el6.text

        table = doc.add_table(2,2 ,style ='Table Grid')
        a = table.cell(0,0).paragraphs[0].add_run("測試內容：美股專區畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['A44'] = "美股行情"
        ws['B44'] = '確認美股專區畫面是否正確'

        if(b1 == "科技龍頭"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：美股專區畫面正確" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C44'] = "pass"
            ws['E44'] = datetime.now()
            ws['E44'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股專區畫面正確")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：美股專區畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C44'] = "wrong"
            ws['E44'] = datetime.now()
            ws['E44'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股專區畫面錯誤")
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()


        el3 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美股專區(延遲)\"]")
        el3.click()
        
        el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美指成分")
        el1.click()
        el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="道瓊指數")
        b2 = el2.text

        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認美指成分畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B45'] = '確認美指成分畫面是否正確'
        if(b2 == "道瓊指數"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：美指成分畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C45'] = "pass"
            ws['E45'] = datetime.now()
            ws['E45'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美指成分畫面無誤")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：美指成分畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C45'] = "wrong"
            ws['E45'] = datetime.now()
            ws['E45'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美指成分畫面錯誤")

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()

        el4 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美指成分(延遲)\"]")
        el4.click()
        el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美股焦點")
        el5.click()
        el6 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美股-科技龍頭股")
        b3 = el6.text

        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認美股焦點畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B46'] = "確認美股焦點畫面是否正確"

        if(b3 == "美股-科技龍頭股"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股焦點畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C46'] = 'pass'
            ws['E46'] = datetime.now()
            ws['E46'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股焦點畫面無誤")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股焦點畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C46'] = 'wrong'
            ws['E46'] = datetime.now()
            ws['E46'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股焦點畫面錯誤")

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()

        sleep(5)
        el4 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美股焦點(延遲)\"]")
        el4.click()
        el8 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美股報價")
        el8.click()
        sleep(5)
        el9 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="鋁")
        b4 = el9.text

        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認美股報價畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B47'] = '確認美股報價畫面是否正確'

        if(b4 == "鋁"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股報價畫面正確" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C47'] = 'pass'
            ws['E47'] = datetime.now()
            ws['E47'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股報價畫面正確")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股報價畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C47'] = 'wrong'
            ws['E47'] = datetime.now()
            ws['E47'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股報價畫面錯誤")

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()

        sleep(3)
        el10 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美股報價(延遲)\"]")
        sleep(3)
        el10.click()
        
        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(31, 600)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.move_to_location(31, 520)
        actions.w3c_actions.pointer_action.release()
        actions.perform()

        
        
        el11 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美股排行")
        sleep(5)
        el11.click()
        sleep(5)
        el12 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="成交量排行")
        b5 = el12.text

        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認美股排行畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B48'] = '確認美股排行畫面是否正確'

        if(b5 == "成交量排行"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股排行畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C48'] = 'pass'
            ws['E48'] = datetime.now()
            ws['E48'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股排行畫面無誤")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股排行畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C48'] = 'wrong'
            ws['E48'] = datetime.now()
            ws['E48'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股排行畫面錯誤")

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()

        el13 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"即時排行(延遲)\"]")
        el13.click()
        el10 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美股報價(延遲)\"]")
        sleep(3)
        el10.click()
        el15 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="美股特別股")
        el15.click()
        sleep(5)
        el16 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="漲跌")
        b6 = el16.text
        ws['B49'] = '確認美股特別股畫面是否正確'
        if(b6 == "漲跌"):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股特別股畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C49'] = 'pass'
            ws['E49'] = datetime.now()
            ws['E49'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股特別股畫面無誤")
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:美股特別股畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C49'] = 'wrong'
            ws['E49'] = datetime.now()
            ws['E49'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print("美股特別股畫面錯誤")

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()
        doc.save('test.docx')
        wb.save('test.xlsx')
        
        el10 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"美股特別股(延遲)\"]")
        sleep(3)
        el10.click()
        
        el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="自選")
        el2.click()
        print("離開美股行情測試")