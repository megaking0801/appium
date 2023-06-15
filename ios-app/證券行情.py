from appium.webdriver.common.appiumby import AppiumBy
from time import sleep
from docx import Document
from appium.webdriver.common.mobileby import MobileBy
from appium.webdriver.common.touch_action import TouchAction
from selenium.webdriver.common.action_chains import ActionChains,ActionBuilder
from selenium.webdriver.common.actions.pointer_input import PointerInput
from selenium.webdriver.common.actions import interaction
from docx import Document
from openpyxl import load_workbook
from docx.shared import Cm,Pt
from datetime import datetime
from openpyxl import Workbook

class Securities:
    def __init__(driver):
        
        doc = Document('test.docx')
        wb = load_workbook('test.xlsx')
        ws = wb.active

        el1 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"自選庫存\"]")
        sleep(3)
        el1.click()

        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(238, 400)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.pause(0.1)
        actions.w3c_actions.pointer_action.release()
        actions.perform()
    

        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(236, 502)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.pause(0.1)
        actions.w3c_actions.pointer_action.release()
        actions.perform()
    
        el3 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="國際金融")
        el3.click()
        el4 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="國際指數")
        a11 = el4.text

        table = doc.add_table(2,2 ,style ='Table Grid')
        a = table.cell(0,0).paragraphs[0].add_run("測試內容：國際金融畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['A38'] = "other"
        ws['B38'] = '確認國際金融畫面是否正確'
        

        if (a11 == '國際指數'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：國際金融畫面正確" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C38'] = "pass"
            ws['E38'] = datetime.now()
            ws['E38'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('國際金融畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：國際金融畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C38'] = "wrong"
            ws['E38'] = datetime.now()
            ws['E38'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('國際金融畫面bad')

        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()
        
        
        sleep(3)
        el1 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"國際金融\"]")
        el1.click()
        el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="熱門排行")
        sleep(3)
        el5.click()
        sleep(3)
        el6 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="上櫃成交量排行")
        sleep(3)
        a1 = el6.text
        
        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面名稱是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        
        ws['B39'] = '確認熱門排行畫面是否正確'
        if (a1 == '上櫃成交量排行'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：熱門排行畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C39'] = "pass"
            ws['E39'] = datetime.now()
            ws['E39'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('熱門排行畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果：熱門排行畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C39'] = "wrong"
            ws['E39'] = datetime.now()
            ws['E39'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('熱門排行畫面錯誤')
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()
        
        
        sleep(3)
        el2 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"熱門排行\"]")
        el2.click()
        el7 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="台股指數")
        el7.click()
        sleep(3)
        el8 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="加權指數")
        a2 = el8.text
        
        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認台股指數畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B40'] = "確認台股指數畫面是否正確"
        
        if (a2 == '加權指數'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:台股指數畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C40'] = 'pass'
            ws['E40'] = datetime.now()
            ws['E40'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('台股指數畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:台股指數畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C40'] = 'wrong'
            ws['E40'] = datetime.now()
            ws['E40'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('台股指數畫面錯誤')
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()

        el2 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"台股指數\"]")
        el2.click()

        actions = ActionChains(driver)
        actions.w3c_actions = ActionBuilder(driver, mouse=PointerInput(interaction.POINTER_TOUCH, "touch"))
        actions.w3c_actions.pointer_action.move_to_location(18, 600)
        actions.w3c_actions.pointer_action.pointer_down()
        actions.w3c_actions.pointer_action.move_to_location(18, 300)
        actions.w3c_actions.pointer_action.release()
        actions.perform()
        el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="類股報價")
        el1.click()
        el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="台股上市")
        a3 = el2.text
        
        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認類股報價畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B41'] = '確認類股報價畫面是否正確'

        if (a3 == '台股上市'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:類股報價畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C41'] = 'pass'
            ws['E41'] = datetime.now()
            ws['E41'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('類股報價畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:類股報價畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C41'] = 'wrong'
            ws['E41'] = datetime.now()
            ws['E41'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('類股報價畫面錯誤')
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()


        el2 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"類股報價\"]")
        el2.click()
        el4 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="概念股報價")
        el4.click()
        el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="台股概念")
        a4 = el5.text
        
        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認概念股報價畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B42'] = '確認概念股報價畫面是否正確'

        if (a4 == '台股概念'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:概念股報價畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C42'] = 'pass'
            ws['E42'] = datetime.now()
            ws['E42'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('概念股報價畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:概念股報價畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C42'] = 'wrong'
            ws['E42'] = datetime.now()
            ws['E42'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('概念股報價畫面錯誤')
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()
        sleep(3)
        el2 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeButton[@name=\"概念股報價\"]")
        el2.click()
        el6 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="法人進出")
        el6.click()
        el7 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="三大法人上市買賣金額")
        a5 = el7.text
       
        table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
        a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認法人進出畫面是否正確" )
        a.font.name='黑體'
        a.font.size=Pt(16)
        ws['B43'] = '確認法人進出畫面是否正確'
        
        if (a5 == '三大法人上市買賣金額'):
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:法人進出畫面畫面無誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C43'] = 'pass'
            ws['E43'] = datetime.now()
            ws['E43'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('法人進出畫面無誤')
        else:
            b = table.cell(0,1).paragraphs[0].add_run("測試結果:法人進出畫面畫面錯誤" )
            b.font.name='黑體'
            b.font.size=Pt(16)
            ws['C43'] = 'pass'
            ws['E43'] = datetime.now()
            ws['E43'].number_format = 'yyyy-mm-dd hh:mm:ss'
            print('法人進出畫面錯誤')
        
        picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
        picture.font.name='黑體'
        picture.font.size=Pt(16)

        driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        run=table.cell(1,0).paragraphs[0].add_run()
        getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
        getPicture.height=Cm(18)
        getPicture.width=Cm(9) 


        other = table.cell(1,1).paragraphs[0].add_run("備註：" )
        other.font.name='黑體'
        other.font.size=Pt(16)

        doc.add_page_break()
        doc.save('test.docx')
        wb.save('test.xlsx')
        sleep(5)



        el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="自選")
        el2.click()
        print("離開證券行情測試")
        
        