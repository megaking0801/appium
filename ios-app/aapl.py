from appium.webdriver.common.appiumby import AppiumBy
from time import sleep
from docx.shared import Cm,Pt
from docx import Document
from openpyxl import load_workbook
from datetime import datetime
from appium.webdriver.common.mobileby import MobileBy

class SearchAapl:
    def __init__(driver):
            
            wb = load_workbook('test.xlsx')
            ws = wb.active
            
            doc = Document('test.docx')
            sleep(3)
            aapl = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="AAPL")
            companyName = aapl.text
            sleep(3)
            aapl.click()
            testName = companyName+" 名稱測試"
            print(testName)
            
			#檢查自選頁面數值
            try:
                doc = Document('test.docx')
                sleep(3)
                
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面數值是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['A35'] = "AAPL"
                elements = driver.find_elements(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText')
                my_list = []
                for element in elements:
                    my_list.append(element.text)
                print(my_list)
                
                a1 = my_list[29],my_list[24],my_list[35],my_list[30]
                a2 = my_list[25],my_list[26],my_list[31],my_list[32]
                a3 = my_list[27],my_list[28],my_list[33],my_list[34]
                a4 = my_list[22],my_list[18],my_list[23],my_list[19]
                print(a1)
                print(a2)
                print(a3)
                print(a4)
                
                ws['B35'] = my_list[29]+my_list[24]
                ws['C35'] = my_list[35]+my_list[30]
                ws['D35'] = my_list[25]+my_list[26]
                ws['E35'] = datetime.now()
                ws['E35'].number_format = 'yyyy-mm-dd hh:mm:ss'
                ws['B36'] = my_list[31]+my_list[32]
                ws['C36'] = my_list[27]+my_list[28]
                ws['D36'] = my_list[33]+my_list[34]
                ws['E36'] = datetime.now()
                ws['E36'].number_format = 'yyyy-mm-dd hh:mm:ss'
                ws['B37'] = my_list[22]+my_list[18]
                ws['C37'] = my_list[23]+my_list[19]
                ws['E37'] = datetime.now()
                ws['E37'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
                b1 = table.cell(0,1).paragraphs[0].add_run(a1)
                b1.add_break()
                b1.font.name='黑體'
                b1.font.size=Pt(14)
                b2 = table.cell(0,1).paragraphs[0].add_run(a2)
                b2.add_break()
                b2.font.name='黑體'
                b2.font.size=Pt(14)
                b3 = table.cell(0,1).paragraphs[0].add_run(a3)
                b3.add_break()
                b3.font.name='黑體'
                b3.font.size=Pt(14)
                b4 = table.cell(0,1).paragraphs[0].add_run(a4)
                b4.font.name='黑體'
                b4.font.size=Pt(14)
        
        
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)
                doc.add_page_break()
                doc.save('test.docx')
            except Exception:
                    print("exception")
			
            ws['A14'] = "AAPL"
			#檢查名稱
            try:
                doc = Document('test.docx')                
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B14'] = '確認自選頁面名稱是否正確'
                
                if(companyName!="AAPL"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：自選頁面名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C14'] = "wrong"
                    ws['E14'] = datetime.now()
                    ws['E14'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("自選頁面名稱錯誤")
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：自選頁面名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C14'] = "pass"
                    ws['E14'] = datetime.now()
                    ws['E14'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("自選頁面名稱無誤")

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                
            except Exception:
                doc = Document('test.docx')
                ws['B14'] = '確認自選頁面名稱是否正確'
                print("exception")
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)

                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
                b.font.name='黑體'
                b.font.size=Pt(16)
            

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break() 
                doc.save('test.docx')
                ws['D14'] = 'Exception'

			#檢查K線說明書
            try:
                doc = Document('test.docx')
                sleep(3)
                
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認K線名稱是否正確及是否有說明書" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B15'] = "確認K線是否有說明書"
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)
                
                sleep(3)
                el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "K線"`]')
                sleep(3)
                el1.click()
                sleep(3)
                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                
                    
                
                a = table.cell(0,1).paragraphs[0].add_run("測試結果：無說明書")
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['C15'] = 'No'
                ws['E15'] = datetime.now()
                ws['E15'].number_format = 'yyyy-mm-dd hh:mm:ss'
                print("測試結果：無說明書")
                doc.add_page_break()
                doc.save('test.docx') 
                
                
            except Exception:
                doc = Document('test.docx')
                ws['B15'] = "確認K線是否有說明書"
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("Exception")
                a.font.name='黑體'
                a.font.size=Pt(16)
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)
                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break()
                doc.save('test.docx')
                
                ws['D15'] = 'Exception'
                ws['E15'] = datetime.now()
                ws['E15'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
                el13 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "X |"`]')
                sleep(3)
                el13.click()
                sleep(3)
            
			
			#KLine Name
            try:
                doc = Document('test.docx')
                KLineName = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="AAPL")
                KLineName = KLineName.text
                print(KLineName)

            
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認K線名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B16'] = '確認K線名稱是否正確'
                
                if(KLineName!="AAPL"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果:K線名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C16'] = 'wrong'
                    ws['E16'] = datetime.now()
                    ws['E16'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("K線名稱錯誤")
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果:K線名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C16'] = 'pass'
                    ws['E16'] = datetime.now()
                    ws['E16'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("K線名稱無誤")

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                
                el13 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "X |"`]')
                sleep(3)
                el13.click()
                sleep(3)
                
            except Exception:
                doc = Document('test.docx')
                ws['B16'] = '確認K線名稱是否正確'
                print("exception")
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認KLine名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)

                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
                b.font.name='黑體'
                b.font.size=Pt(16)
            

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break()
                doc.save('test.docx')
                ws['D16'] = 'Exception'
                ws['E16'] = datetime.now()
                ws['E16'].number_format = 'yyyy-mm-dd hh:mm:ss'



    #買進賣出確認
	        
            el3 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeWindow[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeScrollView/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[3]/XCUIElementTypeOther[3]/XCUIElementTypeButton')
            sleep(3)
            el3.click()
            sleep(3)
            
            #買進
            try: 
                el4 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="CathayIntroduction SBStockOrde")
                sleep(3)
                el4.click()
                sleep(3)
                el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="CathayIntroduction SBStockOrde")
                sleep(3)
                el5.click()
                sleep(3)
                el6 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="CathayIntroduction SBStockOrde")
                sleep(3)
                el6.click()
                sleep(3)
                '''el7 = driver.find_element(by=AppiumBy.XPATH, value="//XCUIElementTypeStaticText[@name=\"暫不申請\"]")
                sleep(3)
                el7.click()'''

                doc = Document('test.docx')
                table = doc.add_table(2,2 ,style ='Table Grid')
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認買進賣出畫面是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B17'] = '確認買進畫面是否正確'
                el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="蘋果")
                sleep(3)
                Name = el5.text

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                
                
                if(Name != "蘋果"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：買進名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    print("買進名稱錯誤")
                    ws['C17'] = 'wrong'
                    ws['E17'] = datetime.now()
                    ws['E17'].number_format = 'yyyy-mm-dd hh:mm:ss'
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：買進名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C17'] = 'pass'
                    ws['E17'] = datetime.now()
                    ws['E17'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("買進名稱無誤")
                
                doc.add_page_break()
                doc.save('test.docx')
                sleep(3)
                el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "下單"`]')
                sleep(3)
                el1.click()

            except Exception:
                doc = Document('test.docx')
                ws['B17'] = '確認買進畫面是否正確'
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                ws['D17'] = 'Exception'
                ws['E17'] = datetime.now()
                ws['E17'].number_format = 'yyyy-mm-dd hh:mm:ss'
            
            
            
            
            #sell
            try:
                doc = Document('test.docx')
                sleep(3)
                el6 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeWindow[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeScrollView/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[3]/XCUIElementTypeOther[5]/XCUIElementTypeButton')
                sleep(3)
                el6.click()
                sleep(3)
                
                el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="蘋果")
                sleep(3)
                Name = el5.text
                ws['B18'] = '確認賣出畫面是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認賣出名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                

                if(Name != "蘋果"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：賣出名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C18'] = 'wrong'
                    ws['E18'] = datetime.now()
                    ws['E18'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("賣出名稱錯誤")
                    
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：賣出名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C18'] = 'pass'
                    ws['E18'] = datetime.now()
                    ws['E18'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("賣出名稱無誤")
                
                doc.add_page_break()
                doc.save('test.docx')
                sleep(3)
                el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "下單"`]')
                sleep(3)
                el1.click()

            except Exception:
                doc = Document('test.docx')
                ws['B18'] = '確認賣出畫面是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid')
                a = table.cell(0,0).paragraphs[0].add_run("Exception" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                ws['D18'] = 'Exception'
                ws['E18'] = datetime.now()
                ws['E18'].number_format = 'yyyy-mm-dd hh:mm:ss'
            
			
            wb.save('test.xlsx')
            sleep(3)
            el2 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "自選1"`]')
            sleep(3)
            el2.click()

