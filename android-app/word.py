from docx import Document
from docx.shared import Cm,Pt
import datetime

dt = datetime.datetime.now()
excelNameTime = str(dt.year) +'年'+str(dt.month)+'月'+str(dt.day)+'日'+str(dt.hour)+'點'+str(dt.minute)+'分'

#cell_new = table1.cell(2,0).merge(table1.cell(2,1)) //合併表格

doc = Document()
#設置標題，儲檔
doc.add_heading('國泰樹精靈測試')
doc.save('test.docx')

class PrintWord:

	def title(driver):
		
		#table = doc.add_table(2,2,style ='Table Grid' )#添加表格

		#company = table.cell(0,0).paragraphs[0].add_run("測試內容：")
		#company.font.name='黑體'
		#company.font.size=Pt(16)

		#result = table.cell(0,1).paragraphs[0].add_run("測試結果：")
		#result.font.name='黑體'
		#result.font.size=Pt(16)

		#picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		#picture.font.name='黑體'
		#picture.font.size=Pt(16)

		#run=table.cell(1,0).paragraphs[0].add_run()
		#getPicture =run.add_picture("/Users/2200290/Downloads/test.png")
		#getPicture.height=Cm(4.4)
		#getPicture.width=Cm(6.2)

		#other = table.cell(1,1).paragraphs[0].add_run("備註：")
		#other.font.name='黑體'
		#other.font.size=Pt(16)

		#換頁
		#doc.add_page_break()

		doc.save('test.docx')

class CompanyName:
	def companyName(driver,company,result,picture,other):
		table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容：" + company)
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果：" + result)
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 


		other = table.cell(1,1).paragraphs[0].add_run("備註：" )
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')	


class TestOddLet:
	def oddLet(driver,companyOddLet,resultOddLet,pictureOddLet,otherOddLet):
		table = doc.add_table(2,2,style ='Table Grid' ) #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容："+companyOddLet+" 零股測試")
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果："+resultOddLet)
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 

		other = table.cell(1,1).paragraphs[0].add_run("備註：")
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')		

class TestKLine:
	def kLine(driver,companyKLine,resultKLine,pictureKLine,otherKLine):
		table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容：" + companyKLine +" K線測試")
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果：" + resultKLine)
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 

		other = table.cell(1,1).paragraphs[0].add_run("備註：")
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')		

#外資、投信...
class CompareNumberList:
	def compareNumberList(driver,company,result,picture,other,testName,buyAndSell):
		table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容：" + company + testName + " " +buyAndSell)
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果：" + result)
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 

		other = table.cell(1,1).paragraphs[0].add_run("備註：")
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')		

class BuyAndSell:
	def buyAndSell(driver,company,result,picture,other,testName):
		table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容：" + company + testName)
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果：" + result)
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 

		other = table.cell(1,1).paragraphs[0].add_run("備註：")
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')	

class Error:
	def error(driver):
		table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

		company = table.cell(0,0).paragraphs[0].add_run("測試內容：")
		company.font.name='黑體'
		company.font.size=Pt(16)

		result = table.cell(0,1).paragraphs[0].add_run("測試結果：" + " 發生錯誤")
		result.font.name='黑體'
		result.font.size=Pt(16)

		picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
		picture.font.name='黑體'
		picture.font.size=Pt(16)

		driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		run=table.cell(1,0).paragraphs[0].add_run()
		getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
		getPicture.height=Cm(18)
		getPicture.width=Cm(9) 


		other = table.cell(1,1).paragraphs[0].add_run("備註：" )
		other.font.name='黑體'
		other.font.size=Pt(16)

		doc.add_page_break()

		doc.save('test.docx')			


