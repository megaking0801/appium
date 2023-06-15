from appium.webdriver.common.appiumby import AppiumBy
from time import sleep
from docx.shared import Cm,Pt
from docx import Document
from openpyxl import Workbook
from datetime import datetime
import openpyxl
from appium.webdriver.common.mobileby import MobileBy

class SearchSystex:
    def __init__(driver):
            doc = Document()
            wb = Workbook()
            ws = wb.active
            
            ws['A1'] = '公司名稱'
            ws['B1'] = '測試內容'
            ws['C1'] = '測試結果'
            ws['D1'] = '備註'
            ws['E1'] = '時間'
            
            doc.add_heading('國泰樹精靈測試')
            doc.save('test.docx')
            
			#點擊公司
            systex = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠 6214")
            companyName = systex.text
            systex.click()
            
            ws['A2'] = "Systex"
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el1.click()
            sleep(3)
            el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el2.click()
            sleep(3)
            el3 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el3.click()
            sleep(3)


            testName = companyName+" 名稱測試"
            print(testName)
            #檢查自選頁面數值
            try:
                doc = Document('test.docx')
                sleep(3)
                ws['A19'] = "Systex"
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面數值是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                
                elements = driver.find_elements(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText')
                my_list = []
                for element in elements:
                    my_list.append(element.text)
                
                a1 = my_list[40],my_list[35],my_list[46],my_list[41]
                a2 = my_list[36],my_list[37],my_list[42],my_list[43]
                a3 = my_list[38],my_list[39],my_list[44],my_list[45]
                a4 = my_list[31],my_list[27],my_list[32],my_list[28]
                print(a1)
                print(a2)
                print(a3)
                print(a4)
                
                ws['B19'] = my_list[40]+my_list[35]
                ws['C19'] = my_list[46]+my_list[41]
                ws['D19'] = my_list[36]+my_list[37]
                ws['E19'] = datetime.now()
                ws['E19'].number_format = 'yyyy-mm-dd hh:mm:ss'
                ws['B20'] = my_list[42]+my_list[43]
                ws['C20'] = my_list[38]+my_list[39]
                ws['D20'] = my_list[44]+my_list[45]
                ws['E20'] = datetime.now()
                ws['E20'].number_format = 'yyyy-mm-dd hh:mm:ss'
                ws['B21'] = my_list[31]+my_list[27]
                ws['C21'] = my_list[32]+my_list[28]
                ws['E21'] = datetime.now()
                ws['E21'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
                b1 = table.cell(0,1).paragraphs[0].add_run(a1)
                b1.add_break()
                b1.font.name='黑體'
                b1.font.size=Pt(14)
                b2 = table.cell(0,1).paragraphs[0].add_run(a2)
                b2.add_break()
                b2.font.name='黑體'
                b2.font.size=Pt(14)
                b3 = table.cell(0,1).paragraphs[0].add_run(a3)
                b3.add_break()
                b3.font.name='黑體'
                b3.font.size=Pt(14)
                b4 = table.cell(0,1).paragraphs[0].add_run(a4)
                b4.font.name='黑體'
                b4.font.size=Pt(14)
        
        
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)
                doc.save('test.docx')
            except Exception:
                print("exception")
			
            a = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeNavigationBar[`name == "TWCTStockOverviewVC"`]/XCUIElementTypeButton[5]')
            sleep(3)
            a.click()
            sleep(3)
			
		#外資
            doc = Document('test.docx')
            el12 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeWindow[1]/XCUIElementTypeOther[2]/XCUIElementTypeOther/XCUIElementTypeTable/XCUIElementTypeCell[9]/XCUIElementTypeOther[1]/XCUIElementTypeOther')
            sleep(3)
            el12.click()
            sleep(30)

            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：外資日期標題是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            ws['B22'] = "外資日期標題是否正確"
            
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="日期")
            sleep(3)
            a = el1.text
            if(a != "日期"):
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：外資日期標題錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                ws['C22'] = "wrong"
                ws['E22'] = datetime.now()
                ws['E22'].number_format = 'yyyy-mm-dd hh:mm:ss'
                print("外資日期標題錯誤")
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：外資日期標題無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("外資日期標題無誤")
                ws['C22'] = "pass"
                ws['E22'] = datetime.now()
                ws['E22'].number_format = 'yyyy-mm-dd hh:mm:ss'
            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()
            doc.save('test.docx')

    #投信
            doc = Document('test.docx')
            el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value=" 投信 ")
            sleep(3)
            el2.click()
            sleep(3)
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：投信日期標題是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            ws['B23'] = "投信日期標題是否正確"
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="日期")
            a = el1.text
            if(a != "日期"):
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：投信日期標題錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("投信日期標題錯誤")
                ws['C23'] = "wrong"
                ws['E23'] = datetime.now()
                ws['E23'].number_format = 'yyyy-mm-dd hh:mm:ss'
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：投信日期標題無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("投信日期標題無誤")
                ws['C23'] = "pass"
                ws['E23'] = datetime.now()
                ws['E23'].number_format = 'yyyy-mm-dd hh:mm:ss'
            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)
    
            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()
            doc.save('test.docx')
    
    #自營
            doc = Document('test.docx')
            el3 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value=" 自營 ")
            sleep(3)
            el3.click()
            sleep(3)
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：自營日期標題是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            ws['B24'] = "自營日期標題是否正確"
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="日期")
            a = el1.text
            if(a != "日期"):
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：自營日期標題錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("自營日期標題錯誤")
                ws['C24'] = "wrong"
                ws['E24'] = datetime.now()
                ws['E24'].number_format = 'yyyy-mm-dd hh:mm:ss'
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：自營日期標題無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("自營日期標題無誤")
                ws['C24'] = "pass"
                ws['E24'] = datetime.now()
                ws['E24'].number_format = 'yyyy-mm-dd hh:mm:ss'
            
            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()
            doc.save('test.docx')

    #融資
            doc = Document('test.docx')
            el4 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value=" 融資 ")
            sleep(3)
            el4.click()
            sleep(3)
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：融資日期標題是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            ws['B25'] = "融資日期標題是否正確"
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="日期")
            a = el1.text
            if(a != "日期"):
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：融資日期標題錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("融資日期標題錯誤")
                ws['C25'] = "wrong"
                ws['E25'] = datetime.now()
                ws['E25'].number_format = 'yyyy-mm-dd hh:mm:ss'
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：融資日期標題無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("融資日期標題無誤")
                ws['C25'] = "pass"
                ws['E25'] = datetime.now()
                ws['E25'].number_format = 'yyyy-mm-dd hh:mm:ss'
            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()
            doc.save('test.docx')
    #融券
            doc = Document('test.docx')
            el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value=" 融券 ")
            sleep(3)
            el5.click()
            sleep(3)
            table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
            a = table.cell(0,0).paragraphs[0].add_run("測試內容：融券日期標題是否正確" )
            a.font.name='黑體'
            a.font.size=Pt(16)
            ws['B26'] = "融券日期標題是否正確"
            el1 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="日期")
            a = el1.text
            if(a != "日期"):
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：融券日期標題錯誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("融券日期標題錯誤")
                ws['C26'] = "wrong"
                ws['E26'] = datetime.now()
                ws['E26'].number_format = 'yyyy-mm-dd hh:mm:ss'
            else:
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：融券日期標題無誤" )
                b.font.name='黑體'
                b.font.size=Pt(16)
                print("融券日期標題無誤")
                ws['C26'] = "pass"
                ws['E26'] = datetime.now()
                ws['E26'].number_format = 'yyyy-mm-dd hh:mm:ss'
            picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
            picture.font.name='黑體'
            picture.font.size=Pt(16)

            driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            run=table.cell(1,0).paragraphs[0].add_run()
            getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
            getPicture.height=Cm(18)
            getPicture.width=Cm(9) 


            other = table.cell(1,1).paragraphs[0].add_run("備註：" )
            other.font.name='黑體'
            other.font.size=Pt(16)

            doc.add_page_break()
            doc.save('test.docx')
			
			
			#檢查名稱
            try:
                doc = Document('test.docx')
                sleep(3)
                
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B2'] = '確認自選頁面名稱是否正確'
                
                if(companyName!="精誠 6214"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：自選頁面名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C2'] = "wrong"
                    ws['E2'] = datetime.now()
                    ws['E2'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("自選頁面名稱錯誤")
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：自選頁面名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C2'] = "pass"
                    ws['E2'] = datetime.now()
                    ws['E2'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("自選頁面名稱無誤")

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)
                ws['C1'] = '備註：'
                doc.add_page_break()
                doc.save('test.docx')
                
            except Exception:
                doc = Document('test.docx')
                ws['B2'] = '確認自選頁面名稱是否正確'
                print("exception")
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認自選頁面名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                
                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
                b.font.name='黑體'
                b.font.size=Pt(16)
            

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break() 
                doc.save('test.docx')
                ws['D2'] = 'Exception'
                sleep(3)
			#檢查零股
            try:
                doc = Document('test.docx')
                sleep(3)
                enterOddLot = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "零股"`]') #零股
                sleep(3)
                enterOddLot.click()

                oddLot = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠 6214")
                
                oddLotName = oddLot.text
                print(oddLotName)

            
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認零股名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B3'] = '確認零股名稱是否正確'
                
                if(oddLotName!="精誠 6214"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：零股名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C3'] = "wrong"
                    ws['E3'] = datetime.now()
                    ws['E3'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("零股名稱錯誤")
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：零股名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C3'] = "pass"
                    ws['E3'] = datetime.now()
                    ws['E3'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("零股名稱無誤")

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                #doc.add_page_break()

                doc.save('test.docx')
            except Exception:
                doc = Document('test.docx')
                print("exception")
                ws['B3'] = '確認零股名稱是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認零股名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)

                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
                b.font.name='黑體'
                b.font.size=Pt(16)
            

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 
                doc.save('test.docx')
                ws['D3'] = "Exception"
        
            doc.add_page_break()
            doc.save('test.docx')
        #離開
            el9 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "零股"`]')
            sleep(3)
            el9.click()
            
			
		#檢查K線  
            try:
                doc = Document('test.docx')
                sleep(3)
                
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認K線名稱是否正確及是否有說明書" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B4'] = "確認K線是否有說明書"
        
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)
                
                
                el2 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "K線"`]')
                sleep(3)
                el2.click()
                sleep(3)
                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                el7 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
                sleep(3)
                el7.click()
                sleep(3)
                el8 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
                sleep(3)
                el8.click()
                sleep(3)
                    
                
                a = table.cell(0,1).paragraphs[0].add_run("測試結果：有說明書")
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['C4'] = 'yes'
                ws['E4'] = datetime.now()
                ws['E4'].number_format = 'yyyy-mm-dd hh:mm:ss'
                doc.add_page_break()
                doc.save('test.docx')    
                
            except Exception:
                doc = Document('test.docx')
                ws['B4'] = "確認K線是否有說明書"
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("Exception")
                a.font.name='黑體'
                a.font.size=Pt(16)
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)
                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break()
                doc.save('test.docx')  
                ws['C4'] = 'No'
                ws['D4'] = 'Exception'
                ws['E4'] = datetime.now()
                ws['E4'].number_format = 'yyyy-mm-dd hh:mm:ss'

            try:
                doc = Document('test.docx')
                KLineName = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="6214 精誠")
                KLineName = KLineName.text
                print(KLineName)

            
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格

                a = table.cell(0,0).paragraphs[0].add_run("測試內容:確認K線名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B5'] = '確認K線名稱是否正確'
                if(KLineName!="6214 精誠"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果:K線名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C5'] = 'wrong'
                    ws['E5'] = datetime.now()
                    ws['E5'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("K線名稱錯誤")
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果:K線名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    ws['C5'] = 'pass'
                    ws['E5'] = datetime.now()
                    ws['E5'].number_format = 'yyyy-mm-dd hh:mm:ss'
                    print("K線名稱無誤")

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                
                el13 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "X |"`]')
                sleep(3)
                el13.click()
                sleep(3)
                
            except Exception:
                doc = Document('test.docx')
                print("exception")
                ws['B5'] = '確認K線名稱是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認KLine名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)

                b = table.cell(0,1).paragraphs[0].add_run("測試結果：" + "exception" )
                b.font.name='黑體'
                b.font.size=Pt(16)
            

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9)
                doc.add_page_break()
                doc.save('test.docx')
                ws['D5'] = 'Exception'
                ws['E5'] = datetime.now()
                ws['E5'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
	#檢查買進、賣出
			
			#買進
            el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeWindow[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeScrollView/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[3]/XCUIElementTypeOther[3]/XCUIElementTypeButton')
            sleep(3)
            el1.click()
            sleep(3)
            el2 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el2.click()
            sleep(3)
            el3 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="MBkResourcesImage Introduction")
            sleep(3)
            el3.click()
            sleep(10)
            el4 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeStaticText[`label == "暫不申請"`]')
            sleep(3)
            el4.click()
            
			
			#buy
            try:
                doc = Document('test.docx')
                table = doc.add_table(2,2 ,style ='Table Grid')
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認買進賣出畫面是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                ws['B6'] = '確認買進畫面是否正確'
                
                el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠")
                Name = el5.text

                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                
                
                if(Name != "精誠"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：買進名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    print("買進名稱錯誤")
                    ws['C6'] = 'wrong'
                    ws['E6'] = datetime.now()
                    ws['E6'].number_format = 'yyyy-mm-dd hh:mm:ss'
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：買進名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    print("買進名稱無誤")
                    ws['C6'] = 'pass'
                    ws['E6'] = datetime.now()
                    ws['E6'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
                doc.add_page_break()
                doc.save('test.docx')
                el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "下單"`]')
                el1.click()

            except Exception:
                doc = Document('test.docx')
                ws['B6'] = '確認買進畫面是否正確'
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                ws['D6'] = 'Exception'
                ws['E6'] = datetime.now()
                ws['E6'].number_format = 'yyyy-mm-dd hh:mm:ss'
            
            
            
            
            #賣出
            try:
                doc = Document('test.docx')
                sleep(3)
                el6 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeWindow[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[1]/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeScrollView/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther/XCUIElementTypeOther[3]/XCUIElementTypeOther[5]/XCUIElementTypeButton')
                sleep(3)
                el6.click()
                sleep(3)
                
                el5 = driver.find_element(by=AppiumBy.ACCESSIBILITY_ID, value="精誠")
                Name = el5.text
                ws['B7'] = '確認賣出畫面是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid') #添加表格
                a = table.cell(0,0).paragraphs[0].add_run("測試內容：確認賣出名稱是否正確" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()

                if(Name != "精誠"):
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：賣出名稱錯誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    print("賣出名稱錯誤")
                    ws['C7'] = 'wrong'
                    ws['E7'] = datetime.now()
                    ws['E7'].number_format = 'yyyy-mm-dd hh:mm:ss'
                else:
                    b = table.cell(0,1).paragraphs[0].add_run("測試結果：賣出名稱無誤" )
                    b.font.name='黑體'
                    b.font.size=Pt(16)
                    print("賣出名稱無誤")
                    ws['C7'] = 'pass'
                    ws['E7'] = datetime.now()
                    ws['E7'].number_format = 'yyyy-mm-dd hh:mm:ss'
                
                doc.save('test.docx')
                el1 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "下單"`]')
                el1.click()

            except Exception:
                doc = Document('test.docx')
                ws['B7'] = '確認賣出畫面是否正確'
                table = doc.add_table(2,2 ,style ='Table Grid')
                a = table.cell(0,0).paragraphs[0].add_run("Exception" )
                a.font.name='黑體'
                a.font.size=Pt(16)
                picture = table.cell(1,0).paragraphs[0].add_run("畫面：")
                picture.font.name='黑體'
                picture.font.size=Pt(16)

                driver.get_screenshot_as_file("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                run=table.cell(1,0).paragraphs[0].add_run()
                getPicture =run.add_picture("/Users/21ms20r/Downloads/app-main/app/testPhoto.png")
                getPicture.height=Cm(18)
                getPicture.width=Cm(9) 


                other = table.cell(1,1).paragraphs[0].add_run("備註：" )
                other.font.name='黑體'
                other.font.size=Pt(16)

                doc.add_page_break()
                doc.save('test.docx')
                ws['D7'] = 'Exception'
                ws['E7'] = datetime.now()
                ws['E7'].number_format = 'yyyy-mm-dd hh:mm:ss'

			
			#離開此公司畫面
            wb.save('test.xlsx')
            el2 = driver.find_element(by=MobileBy.IOS_CLASS_CHAIN, value='**/XCUIElementTypeButton[`label == "自選1"`]')
            el2.click()
            print("離開"+ testName)